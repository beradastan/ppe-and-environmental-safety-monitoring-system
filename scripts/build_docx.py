from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Sayfa düzeni
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(3.0)
sec.left_margin   = Cm(2.75)
sec.bottom_margin = Cm(2.0)
sec.right_margin  = Cm(2.75)

FONT = "Palatino Linotype"

def fmt_run(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic

def fmt_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=18):
    pf = p.paragraph_format
    pf.alignment         = align
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(line)

def special_title(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)
    fmt_run(p.add_run(text), size=12, bold=True)

def ch1(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, before=18, after=6)
    fmt_run(p.add_run(text.upper()), size=14, bold=True)

def ch2(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, before=12, after=6)
    fmt_run(p.add_run(text), size=12, bold=True)

def ch3(doc, text):
    p = doc.add_paragraph()
    fmt_para(p, before=9, after=6)
    fmt_run(p.add_run(text), size=12, bold=True)

def body(doc, text="[İÇERİK BURAYA EKLENECEK]"):
    p = doc.add_paragraph()
    fmt_para(p)
    fmt_run(p.add_run(text), size=12, italic=True)

def pb(doc):
    doc.add_page_break()

# ─── ÖN SAYFALAR ─────────────────────────────────────────────────────────────
special_title(doc, "ÖZET")
body(doc)
pb(doc)
special_title(doc, "ABSTRACT")
body(doc)
pb(doc)
special_title(doc, "TEŞEKKÜR")
body(doc, "[Kişisel teşekkür metni buraya eklenecektir.]")
pb(doc)
special_title(doc, "İÇİNDEKİLER")
body(doc, "[Word'de otomatik içindekiler tablosu eklenecektir.]")
pb(doc)
special_title(doc, "SİMGELER VE KISALTMALAR")
body(doc)

# ─── 1. GİRİŞ ────────────────────────────────────────────────────────────────
pb(doc)
ch1(doc, "1. Giriş")
ch2(doc, "1.1. Problem Tanımı, Amaç ve Hedefler")
body(doc)
ch2(doc, "1.2. Nesne Tespiti ve Derin Öğrenme")
body(doc)
ch2(doc, "1.3. Çok Nesne Takibi")
body(doc)
ch2(doc, "1.4. Büyük Dil Modelleri ve Raporlama")
body(doc)
ch2(doc, "1.5. Bu Çalışmanın Kapsamı ve Katkıları")
body(doc)

# ─── 2. LİTERATÜR ÖZETİ ─────────────────────────────────────────────────────
pb(doc)
ch1(doc, "2. Literatür Özeti")
ch2(doc, "2.1. Nesne Tespiti ve YOLO Mimarisi")
body(doc)
ch2(doc, "2.2. KKD Tespiti Uygulamaları")
body(doc)
ch2(doc, "2.3. Çok Nesne Takip Algoritmaları")
body(doc)
ch2(doc, "2.4. Büyük Dil Modelleri")
body(doc)

# ─── 3. YÖNTEM ───────────────────────────────────────────────────────────────
pb(doc)
ch1(doc, "3. Yöntem")

ch2(doc, "3.1. Araştırma Deseni ve Sistem Mimarisi")
ch3(doc, "3.1.1. Genel sistem tasarımı ve bileşenler")
body(doc)
ch3(doc, "3.1.2. Kullanılan teknolojiler ve araçlar")
body(doc)
ch3(doc, "3.1.3. Konfigürasyon yönetimi")
body(doc)
ch3(doc, "3.1.4. Bileşenler arası iletişim ve veri akışı")
body(doc)

ch2(doc, "3.2. Veri Setleri ve Model Eğitimi")
ch3(doc, "3.2.1. Veri seti seçim kriterleri")
body(doc)
ch3(doc, "3.2.2. Kullanılan veri setleri")
body(doc)
ch3(doc, "3.2.3. Veri ön işleme ve crop veri seti üretimi")
body(doc)
ch3(doc, "3.2.4. Model eğitimi")
body(doc)
ch3(doc, "3.2.5. Model değerlendirme metrikleri")
body(doc)
ch3(doc, "3.2.6. Model karşılaştırmaları ve nihai seçim kararları")
body(doc)

ch2(doc, "3.3. Detection Pipeline")
ch3(doc, "3.3.1. Pipeline genel akışı ve kaynak yönetimi")
body(doc)
ch3(doc, "3.3.2. Kişi tespiti ve ByteTrack takibi")
body(doc)
ch3(doc, "3.3.3. Stabil kişi kimliği — TrackReattacher")
body(doc)
ch3(doc, "3.3.4. Crop-based PPE detection modu")
body(doc)
ch3(doc, "3.3.5. Scene-based PPE detection modu")
body(doc)
ch3(doc, "3.3.6. İki modun tasarım farkları ve trade-off'ları")
body(doc)
ch3(doc, "3.3.7. Yangın ve duman tespiti")
body(doc)
ch3(doc, "3.3.8. Temporal voting (zamansal oylama)")
body(doc)
ch3(doc, "3.3.9. Kamera izleme sistemi")
body(doc)

ch2(doc, "3.4. Olay Yönetimi (Event State Machine)")
ch3(doc, "3.4.1. Motivasyon ve durum geçiş diyagramı")
body(doc)
ch3(doc, "3.4.2. Durum geçiş koşulları")
body(doc)
ch3(doc, "3.4.3. Yangın debounce")
body(doc)
ch3(doc, "3.4.4. Per-kişi ihlal takibi")
body(doc)
ch3(doc, "3.4.5. Event signature, repeat_count ve duration_sec")
body(doc)

ch2(doc, "3.5. Backend Sistemi")
ch3(doc, "3.5.1. Flask uygulama mimarisi")
body(doc)
ch3(doc, "3.5.2. REST API tasarımı")
body(doc)
ch3(doc, "3.5.3. Socket.IO gerçek zamanlı eventler")
body(doc)
ch3(doc, "3.5.4. Veritabanı katmanı")
body(doc)
ch3(doc, "3.5.5. Dosya sistemi izleyici (Watcher)")
body(doc)
ch3(doc, "3.5.6. Pipeline yönetimi")
body(doc)
ch3(doc, "3.5.7. Güvenlik ve doğrulama")
body(doc)

ch2(doc, "3.6. Veritabanı Tasarımı")
ch3(doc, "3.6.1. Tablo yapıları")
body(doc)
ch3(doc, "3.6.2. JSONB kullanımı")
body(doc)
ch3(doc, "3.6.3. İndeks stratejisi ve idempotent schema")
body(doc)
ch3(doc, "3.6.4. false_positive soft-delete yaklaşımı")
body(doc)

ch2(doc, "3.7. LLM Entegrasyonu")
ch3(doc, "3.7.1. Ollama ile yerel LLM çalıştırma ve model seçimi")
body(doc)
ch3(doc, "3.7.2. SafetyReportAgent tasarımı")
body(doc)
ch3(doc, "3.7.3. Rapor analitik servisleri")
body(doc)
ch3(doc, "3.7.4. Periyodik rapor üretimi")
body(doc)

ch2(doc, "3.8. Frontend Sistemi")
ch3(doc, "3.8.1. React + Vite uygulama yapısı")
body(doc)
ch3(doc, "3.8.2. Sayfalar")
body(doc)
ch3(doc, "3.8.3. Gerçek zamanlı güncelleme ve kamera durum banner'ı")
body(doc)
ch3(doc, "3.8.4. Tema sistemi")
body(doc)
ch3(doc, "3.8.5. Export sistemi")
body(doc)

# ─── 4. BULGULAR VE DEĞERLENDİRME ───────────────────────────────────────────
pb(doc)
ch1(doc, "4. Bulgular ve Değerlendirme")

ch2(doc, "4.1. Test Ortamı")
body(doc)

ch2(doc, "4.2. Model Performans Sonuçları")
ch3(doc, "4.2.1. Ortak modeller — kişi ve yangın/duman tespiti")
body(doc)
ch3(doc, "4.2.2. Crop-based PPE modelleri")
body(doc)
ch3(doc, "4.2.3. Scene-based PPE modelleri")
body(doc)

ch2(doc, "4.3. Crop-Based ve Scene-Based Modların Karşılaştırması")
ch3(doc, "4.3.1. Baret tespitinde karşılaştırma")
body(doc)
ch3(doc, "4.3.2. Yelek tespitinde karşılaştırma")
body(doc)
ch3(doc, "4.3.3. Maske tespitinde karşılaştırma")
body(doc)
ch3(doc, "4.3.4. FPS ve GPU bellek kullanımı")
body(doc)
ch3(doc, "4.3.5. Oklüzyon ve kalabalık senaryolarında performans")
body(doc)
ch3(doc, "4.3.6. Genel değerlendirme")
body(doc)

ch2(doc, "4.4. Sistem Performansı")
ch3(doc, "4.4.1. Uçtan uca işlem hızı")
body(doc)
ch3(doc, "4.4.2. TrackReattacher etkinliği")
body(doc)
ch3(doc, "4.4.3. PPE çıkarım sıklığı optimizasyonu (PPE_INFER_EVERY)")
body(doc)
ch3(doc, "4.4.4. PPE tespit confidence eşiği optimizasyonu")
body(doc)
ch3(doc, "4.4.5. Temporal voting penceresi optimizasyonu (temporal_window)")
body(doc)
ch3(doc, "4.4.6. İçerme kesri eşiği optimizasyonu (INSIDE_FRAC_THR)")
body(doc)
ch3(doc, "4.4.7. Scene modu confidence eşiği optimizasyonu")
body(doc)
ch3(doc, "4.4.8. Scene modu temporal voting penceresi optimizasyonu")
body(doc)
ch3(doc, "4.4.9. Crop ve scene modlarının karşılaştırılması")
body(doc)

ch2(doc, "4.5. Ground Truth Değerlendirmesi")
body(doc)

ch2(doc, "4.6. LLM Rapor Kalitesi Değerlendirmesi")
ch3(doc, "4.6.1. Per-alarm LLM raporu")
body(doc)
ch3(doc, "4.6.2. Periyodik rapor")
body(doc)
ch3(doc, "4.6.3. Değerlendirme")
body(doc)
ch3(doc, "4.6.4. Sınırlamalar")
body(doc)

ch2(doc, "4.7. Literatürle Karşılaştırma")
ch3(doc, "4.7.1. Sistem özellikleri karşılaştırması")
body(doc)
ch3(doc, "4.7.2. Model metrikleri karşılaştırması")
body(doc)

# ─── 5. SONUÇ ────────────────────────────────────────────────────────────────
pb(doc)
ch1(doc, "5. Sonuç")
body(doc)

# ─── KAYNAKLAR ───────────────────────────────────────────────────────────────
pb(doc)
special_title(doc, "KAYNAKLAR")
body(doc)

# ─── EKLER ───────────────────────────────────────────────────────────────────
pb(doc)
special_title(doc, "EKLER")

ch2(doc, "EK-1: Fizibilite Raporu")
ch3(doc, "1. Proje tanımı ve kapsam")
ch3(doc, "2. Teknik fizibilite")
ch3(doc, "3. Operasyonel fizibilite")
ch3(doc, "4. Süre fizibilitesi")
ch3(doc, "5. Maliyet fizibilitesi")
ch3(doc, "6. Risk analizi")
ch3(doc, "7. Sonuç")
body(doc)

ch2(doc, "EK-2: Sistem Test Raporu")
ch3(doc, "1. Birim test sonuçları")
ch3(doc, "2. Entegrasyon test sonuçları")
ch3(doc, "3. Performans test sonuçları")
ch3(doc, "4. Bilinen sınırlamalar")
body(doc)

ch2(doc, "EK-3: config.yaml Tam İçeriği")
body(doc)

ch2(doc, "EK-4: Veritabanı Şeması (schema.sql)")
body(doc)

ch2(doc, "EK-5: API Endpoint Referans Tablosu")
body(doc)

ch2(doc, "EK-6: Kesinleşmiş Model Dosyaları ve Seçim Kriterleri")
body(doc)

# ─── ÖZGEÇMİŞ ────────────────────────────────────────────────────────────────
pb(doc)
special_title(doc, "ÖZGEÇMİŞ")
body(doc, "[Kişisel özgeçmiş buraya eklenecektir.]")

doc.save("C:/Users/berat/Desktop/BeraDaştan_TezTaslak.docx")
print("OK")
